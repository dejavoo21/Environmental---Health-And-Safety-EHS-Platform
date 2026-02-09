import re
from pathlib import Path

from docx import Document


SECTION_SOURCES = [
    ("Introduction & Scope", "docs/BRD_EHS_PORTAL_PHASE1.md"),
    ("Business Requirements (Phase 1)", "docs/BRD_EHS_PORTAL_PHASE1.md"),
    ("Architecture Overview", "docs/ARCHITECTURE_PHASE1.md"),
    ("Data Model (ERD)", "docs/DATA_MODEL_PHASE1.md"),
    ("Key Workflows", "docs/WORKFLOWS_PHASE1.md"),
    ("API Summary", "docs/API_SPEC_PHASE1.md"),
    ("Test Strategy (Phase 1)", "docs/TEST_STRATEGY_PHASE1.md"),
]

DIAGRAM_SOURCES = [
    "docs/ARCHITECTURE_PHASE1.md",
    "docs/DATA_MODEL_PHASE1.md",
    "docs/WORKFLOWS_PHASE1.md",
]


def find_mermaid_blocks(text):
    pattern = re.compile(r"```mermaid(.*?)```", re.DOTALL)
    return pattern.findall(text)


def main():
    script_dir = Path(__file__).resolve().parent
    root_dir = script_dir.parent
    output_path = root_dir / "EHS_Phase1_Design.docx"

    doc = Document()
    doc.add_heading("EHS Portal - Phase 1 Design", level=0)

    for heading, source in SECTION_SOURCES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(f"See {source} for full details.")

        if source in DIAGRAM_SOURCES:
            source_path = root_dir / source
            if source_path.exists():
                content = source_path.read_text(encoding="utf-8")
                mermaid_blocks = find_mermaid_blocks(content)
                for index, _block in enumerate(mermaid_blocks, start=1):
                    doc.add_paragraph(
                        f"Insert exported diagram from {source} (diagram {index}) "
                        "PNG/SVG from Mermaid here."
                    )
            else:
                doc.add_paragraph(
                    f"Insert exported diagrams from {source} (PNG/SVG from Mermaid) here."
                )

    doc.save(output_path)
    print(f"Generated {output_path}")


if __name__ == "__main__":
    main()
